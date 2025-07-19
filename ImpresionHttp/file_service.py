import os
import tempfile
import logging
from typing import Dict, Any, Optional, Tuple
from datetime import datetime
import PyPDF2
from docx import Document
from PIL import Image
import io
import mimetypes

class FileService:
    """Servicio para manejar archivos PDF y Word"""
    
    supported_types = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.txt': 'text/plain',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png'
    }

    @staticmethod
    def validate_file(file_path: str):
        _, ext = os.path.splitext(file_path.lower())
        mime_type = FileService.supported_types.get(ext)
        if not mime_type:
            return False, f"Tipo de archivo no soportado: {ext}", ""
        file_size = os.path.getsize(file_path)
        if file_size > 50 * 1024 * 1024:
            return False, "Archivo demasiado grande (máximo 50MB)", ""
        return True, ext[1:], mime_type
    
    def process_file_for_printing(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Procesar archivo para impresión"""
        try:
            if file_type == 'pdf':
                return self.extract_text_from_pdf(file_path)
            elif file_type in ['docx', 'doc']:
                return self.extract_text_from_docx(file_path)
            elif file_type in ['png', 'jpg', 'jpeg']:
                return self.process_image_for_printing(file_path)
            elif file_type == 'txt':
                return self.process_text_file(file_path)
            else:
                return {
                    'success': False,
                    'error': f'Tipo de archivo no soportado: {file_type}'
                }
                
        except Exception as e:
            self.logger.error(f"Error procesando archivo {file_path}: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def process_image_for_printing(self, file_path: str) -> Dict[str, Any]:
        """Procesar imagen para impresión"""
        try:
            with Image.open(file_path) as img:
                # Convertir a RGB si es necesario
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Guardar como PDF temporal
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                    img.save(temp_pdf.name, 'PDF', resolution=300.0)
                    temp_pdf_path = temp_pdf.name
                
                return {
                    'success': True,
                    'file_type': 'image',
                    'original_format': img.format,
                    'size': img.size,
                    'temp_pdf_path': temp_pdf_path,
                    'print_ready': True
                }
                
        except Exception as e:
            self.logger.error(f"Error procesando imagen {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'file_type': 'image'
            }
    
    def process_text_file(self, file_path: str) -> Dict[str, Any]:
        """Procesar archivo de texto"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            return {
                'success': True,
                'file_type': 'text',
                'content': content,
                'lines': len(content.split('\n')),
                'print_ready': True
            }
            
        except Exception as e:
            self.logger.error(f"Error procesando archivo de texto {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'file_type': 'text'
            }
    
    def create_print_job_data(self, file_path: str, file_type: str, 
                            copies: int = 1, priority: int = 2) -> Dict[str, Any]:
        """Crear datos de trabajo de impresión desde archivo"""
        try:
            # Procesar archivo
            processed_data = self.process_file_for_printing(file_path, file_type)
            
            if not processed_data['success']:
                return processed_data
            
            # Crear contenido para impresión
            if file_type == 'pdf':
                content = f"PDF - {processed_data['page_count']} páginas"
                print_data = {
                    'type': 'pdf',
                    'file_path': file_path,
                    'page_count': processed_data['page_count'],
                    'text_content': processed_data['text_content']
                }
            elif file_type in ['docx', 'doc']:
                content = f"Documento Word - {len(processed_data['text_content'])} párrafos"
                print_data = {
                    'type': 'document',
                    'file_path': file_path,
                    'paragraphs': processed_data['paragraphs'],
                    'tables': processed_data['tables'],
                    'text_content': processed_data['text_content']
                }
            elif file_type in ['png', 'jpg', 'jpeg']:
                content = f"Imagen - {processed_data['size'][0]}x{processed_data['size'][1]}"
                print_data = {
                    'type': 'image',
                    'file_path': file_path,
                    'temp_pdf_path': processed_data.get('temp_pdf_path'),
                    'size': processed_data['size']
                }
            else:
                content = f"Archivo {file_type.upper()}"
                print_data = {
                    'type': 'file',
                    'file_path': file_path,
                    'file_type': file_type
                }
            
            return {
                'success': True,
                'content': content,
                'print_data': print_data,
                'copies': copies,
                'priority': priority,
                'file_info': {
                    'original_path': file_path,
                    'file_type': file_type,
                    'file_size': os.path.getsize(file_path),
                    'processed_at': datetime.now().isoformat()
                }
            }
            
        except Exception as e:
            self.logger.error(f"Error creando datos de impresión para {file_path}: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def cleanup_temp_files(self, temp_files: list):
        """Limpiar archivos temporales"""
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
            except Exception as e:
                self.logger.warning(f"No se pudo eliminar archivo temporal {temp_file}: {e}") 

    @staticmethod
    def extract_text_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extraer texto de archivo PDF"""
        try:
            text_content = []
            page_count = 0
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        text_content.append({
                            'page': page_num + 1,
                            'content': text.strip()
                        })
            
            return {
                'success': True,
                'file_type': 'pdf',
                'page_count': page_count,
                'text_content': text_content,
                'total_pages': page_count
            }
            
        except Exception as e:
            self.logger.error(f"Error extrayendo texto de PDF {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'file_type': 'pdf'
            }
    
    @staticmethod
    def extract_text_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extraer texto de archivo Word (.docx)"""
        try:
            doc = Document(file_path)
            text_content = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text.strip())
            
            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return {
                'success': True,
                'file_type': 'docx',
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'text_content': text_content
            }
            
        except Exception as e:
            self.logger.error(f"Error extrayendo texto de DOCX {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'file_type': 'docx'
            }
    
    @staticmethod
    def process_text_file(self, file_path: str) -> Dict[str, Any]:
        """Procesar archivo de texto"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            return {
                'success': True,
                'file_type': 'text',
                'content': content,
                'lines': len(content.split('\n')),
                'print_ready': True
            }
            
        except Exception as e:
            self.logger.error(f"Error procesando archivo de texto {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'file_type': 'text'
            }
    
    @staticmethod
    def process_image_for_printing(self, file_path: str) -> Dict[str, Any]:
        """Procesar imagen para impresión"""
        try:
            with Image.open(file_path) as img:
                # Convertir a RGB si es necesario
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Guardar como PDF temporal
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                    img.save(temp_pdf.name, 'PDF', resolution=300.0)
                    temp_pdf_path = temp_pdf.name
                
                return {
                    'success': True,
                    'file_type': 'image',
                    'original_format': img.format,
                    'size': img.size,
                    'temp_pdf_path': temp_pdf_path,
                    'print_ready': True
                }
                
        except Exception as e:
            self.logger.error(f"Error procesando imagen {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'file_type': 'image'
            } 